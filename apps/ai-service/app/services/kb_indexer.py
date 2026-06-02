"""KB 向量化执行器：解析文档 → chunk → embed → 写 kb_embeddings。

输入由 Go 后端通过 POST /api/v1/kb/{kb_id}/files/{file_id}/index 推送，
请求体 ``KBIndexRequest`` 含原始字节（base64 解码后传入此模块）+ MIME 提示。

支持的文件类型（Phase 2 扩展）：
    text/plain                          → utf-8 解码
    text/markdown                       → utf-8 解码
    text/html                           → trafilatura.extract（fallback 正则去标签）
    application/json                    → 格式化 JSON 字符串
    text/csv                            → 转 markdown 表格（前 1000 行）
    application/pdf                     → pypdf 提取 page-by-page
    application/vnd.openxmlformats-...  → python-docx 段落抽取
    其他                                → ValueError → status=FAILED

实现复用：
    * chunker.split(...)  → 与 post embeddings 同一切片器
    * LlmRouter.embed(...) → 与 post embeddings 同一嵌入路径
    * pgvector 写入语义与 vector_store.upsert_post_embedding 对齐
"""
from __future__ import annotations

import asyncio
import json as _json
import logging
import re
import time
from dataclasses import dataclass

from fastapi import HTTPException

from app.services.chunker import Chunk, split as chunk_split
from app.services.llm_router import LlmRouter

logger = logging.getLogger("ai-service")


# ---------- 文档解析 ----------

_HTML_TAG_RE = re.compile(r"<[^>]+>")
_HTML_SCRIPT_RE = re.compile(r"<(script|style)[^>]*>.*?</\1>", re.IGNORECASE | re.DOTALL)
_WS_RE = re.compile(r"\s+")


def _try_trafilatura(html: str) -> str | None:
    """尝试用 trafilatura.extract 抽取 HTML 主体。失败返回 None 走 fallback。"""
    try:
        import trafilatura  # type: ignore

        extracted = trafilatura.extract(html, favor_recall=True) or ""
        text = extracted.strip()
        return text or None
    except Exception:
        return None


def _strip_html(html: str) -> str:
    """无 trafilatura 时的 fallback：去 script/style + 标签 + 折叠空白。"""
    no_script = _HTML_SCRIPT_RE.sub(" ", html)
    no_tags = _HTML_TAG_RE.sub(" ", no_script)
    collapsed = _WS_RE.sub(" ", no_tags)
    return collapsed.strip()


def extract_pdf_text_pages(content: bytes) -> list[str]:
    """从 PDF 字节流抽取逐页文本。

    依赖 pypdf；缺失时抛 RuntimeError 让调用方写 FAILED。返回值保留空页占位，
    Atlas PDF carrier 需要页码与 char offset 的稳定映射。
    """
    try:
        import pypdf  # type: ignore
    except ImportError as exc:
        raise RuntimeError(f"pypdf 未安装，无法解析 PDF: {exc}") from exc
    import io as _io
    reader = pypdf.PdfReader(_io.BytesIO(content))
    parts: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            txt = page.extract_text() or ""
        except Exception as exc:
            logger.warning("kb_parse.pdf_page_failed", extra={"data": {"page": i, "error": str(exc)[:200]}})
            txt = ""
        parts.append(txt.strip())
    return parts


def _parse_pdf(content: bytes) -> str:
    """从 pdf 字节流抽取所有页面文本，页与页之间用双换行分隔。"""
    return "\n\n".join(page for page in extract_pdf_text_pages(content) if page)


def _parse_docx(content: bytes) -> str:
    """从 docx 字节流抽取段落 + 表格行文本。

    依赖 python-docx；缺失时抛 RuntimeError。
    """
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:
        raise RuntimeError(f"python-docx 未安装，无法解析 DOCX: {exc}") from exc
    import io as _io
    doc = Document(_io.BytesIO(content))
    parts: list[str] = []
    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def _parse_csv(content: bytes) -> str:
    """把 csv 转成 markdown 表格（前 1000 行）。

    标准库 csv 解析，保证可移植不依赖 pandas。
    """
    import csv as _csv
    import io as _io
    text = content.decode("utf-8-sig", errors="replace")
    reader = _csv.reader(_io.StringIO(text))
    rows: list[list[str]] = []
    for i, row in enumerate(reader):
        if i >= 1000:
            rows.append(["…（超过 1000 行已截断）"])
            break
        rows.append([(c or "").strip() for c in row])
    if not rows:
        return ""
    parts: list[str] = []
    header = rows[0]
    parts.append("| " + " | ".join(header) + " |")
    parts.append("| " + " | ".join(["---"] * len(header)) + " |")
    for r in rows[1:]:
        if len(r) < len(header):
            r = r + [""] * (len(header) - len(r))
        elif len(r) > len(header):
            r = r[: len(header)]
        parts.append("| " + " | ".join(r) + " |")
    return "\n".join(parts)


def parse_bytes_to_text(content: bytes, mime_type: str | None, filename: str | None) -> str:
    """把上传的字节流转成纯文本。

    任何无法识别的类型直接抛 ValueError，由调用方写 FAILED。
    解析器缺失（pypdf/python-docx）抛 RuntimeError，同样会被 FAILED 记录。
    """
    if not content:
        return ""
    mime = (mime_type or "").lower().split(";", 1)[0].strip()
    name = (filename or "").lower()

    # 1. PDF（按扩展名优先，再回退 mime —— 部分上传 mime 是 application/octet-stream）
    if name.endswith(".pdf") or mime == "application/pdf":
        return _parse_pdf(content)

    # 2. DOCX —— OOXML 格式；legacy .doc (application/msword) 不走这里
    # （python-docx 仅支持 OOXML；review chatgpt-codex P2：上游已从 KB
    # 白名单移除 application/msword，indexer 也保持一致拒绝）。
    if name.endswith(".docx") or mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return _parse_docx(content)
    if name.endswith(".doc") or mime == "application/msword":
        raise ValueError("legacy .doc 格式不被支持，请用 Word 另存为 .docx 后重新上传")

    # 3. CSV
    if name.endswith(".csv") or mime == "text/csv":
        return _parse_csv(content)

    # 4. text/* 直接 utf-8 decode（带 BOM 兼容）
    if mime.startswith("text/") and not mime.startswith("text/html"):
        return content.decode("utf-8-sig", errors="replace")

    # 5. Markdown by mime or extension
    if mime in {"text/markdown", "text/x-markdown"} or name.endswith((".md", ".markdown", ".mdown")):
        return content.decode("utf-8-sig", errors="replace")

    # 6. JSON
    if mime == "application/json" or name.endswith(".json"):
        try:
            data = _json.loads(content.decode("utf-8-sig", errors="replace"))
            return _json.dumps(data, ensure_ascii=False, indent=2)
        except Exception:
            return content.decode("utf-8-sig", errors="replace")

    # 7. HTML
    if mime == "text/html" or name.endswith((".html", ".htm")):
        html = content.decode("utf-8-sig", errors="replace")
        extracted = _try_trafilatura(html) or _strip_html(html)
        return extracted

    # 8. plain bytes 兜底：如果是文本类（字节中可打印率 > 90%），按 utf-8 decode
    try:
        decoded = content.decode("utf-8")
        printable_ratio = sum(1 for ch in decoded if ch.isprintable() or ch in "\n\r\t") / max(len(decoded), 1)
        if printable_ratio >= 0.9:
            return decoded
    except UnicodeDecodeError:
        pass

    raise ValueError(
        f"不支持的文件类型 mime={mime!r} filename={filename!r}。"
        "支持 txt/md/html/json/csv/pdf/docx。如需其他类型请在 kb_indexer.py 扩展。"
    )


# ---------- KB profile 读取 ----------


@dataclass
class KBActiveProfile:
    id: int
    kb_id: int
    code: str
    model_id: str
    chunker_kind: str
    chunk_size_tokens: int
    chunk_overlap_tokens: int
    top_k: int
    score_threshold: float
    status: str


async def fetch_kb_active_profile(pool, kb_id: int) -> KBActiveProfile:
    """返回 kb_id 上的 active profile；找不到抛 503。"""
    async with pool.acquire() as conn:
        row = await conn.fetchrow(
            """
            SELECT id, kb_id, code, model_id, chunker_kind, chunk_size_tokens,
                   chunk_overlap_tokens, top_k, score_threshold, status
            FROM kb_profiles WHERE kb_id = $1 AND status = 'active' LIMIT 1
            """,
            kb_id,
        )
    if not row:
        raise HTTPException(
            status_code=503,
            detail=f"知识库 {kb_id} 当前没有 active profile，请先在 admin UI 创建并激活",
        )
    return KBActiveProfile(**dict(row))


async def fetch_kb_profile_by_id(pool, profile_id: int) -> KBActiveProfile:
    """按 id 返回任意 profile（含 shadow / deprecated）。蓝绿激活使用。"""
    async with pool.acquire() as conn:
        row = await conn.fetchrow(
            """
            SELECT id, kb_id, code, model_id, chunker_kind, chunk_size_tokens,
                   chunk_overlap_tokens, top_k, score_threshold, status
            FROM kb_profiles WHERE id = $1
            """,
            profile_id,
        )
    if not row:
        raise HTTPException(status_code=404, detail=f"kb profile {profile_id} 不存在")
    return KBActiveProfile(**dict(row))


# ---------- 主流程：解析 → 切片 → embed → 写表 ----------


@dataclass
class KBIndexOutcome:
    kb_file_id: int
    profile_id: int
    chunk_count: int
    doc_chars: int
    doc_tokens: int
    status: str  # 'SUCCEEDED' | 'FAILED'
    error: str = ""


class KBIndexerService:
    """KB 向量化执行器。"""

    def __init__(self, pool, llm: LlmRouter, chunk_concurrency: int = 5) -> None:
        self.pool = pool
        self.llm = llm
        self._chunk_concurrency = chunk_concurrency

    async def vectorize(
        self,
        *,
        kb_id: int,
        kb_file_id: int,
        content_bytes: bytes,
        mime_type: str | None,
        filename: str | None,
        target_profile_id: int | None = None,
        target_status: str = "active",
    ) -> KBIndexOutcome:
        """主入口：解析 → 切片 → embed → 单事务写 kb_embeddings + UPDATE 状态。

        target_profile_id：可选；指定后用该 profile 而非 active（蓝绿 reindex 使用）。
        target_status：写入新行的 status 值；蓝绿走 'shadow'，普通走 'active'。
        """
        try:
            text = parse_bytes_to_text(content_bytes, mime_type, filename)
        except (ValueError, RuntimeError) as exc:
            return KBIndexOutcome(kb_file_id, 0, 0, 0, 0, "FAILED", str(exc))

        # 先解析 profile —— 即使是空文档也要知道 profile_id 才能正确 scope _write_empty
        # （review chatgpt-codex P1 修复：之前空文档走全表 DELETE 会清掉蓝绿 active 行）
        if target_profile_id is not None:
            profile = await fetch_kb_profile_by_id(self.pool, target_profile_id)
            # SECURITY (review chatgpt-codex P2)：profile 必须属于本 KB
            # 否则 attacker 可借 internal token 用 KB B 的 profile_id 索引 KB A
            # 的文件 → kb_embeddings 的 (kb_id, profile_id) 出现错配组合。
            if profile.kb_id != kb_id:
                return KBIndexOutcome(
                    kb_file_id, 0, 0, 0, 0, "FAILED",
                    f"profile {target_profile_id} 不属于 kb {kb_id}",
                )
        else:
            profile = await fetch_kb_active_profile(self.pool, kb_id)

        doc_chars = len(text)
        if doc_chars == 0:
            return await self._write_empty(kb_id, kb_file_id, profile.id)

        chunks: list[Chunk] = chunk_split(
            text,
            chunker_kind=profile.chunker_kind,
            chunk_size_tokens=profile.chunk_size_tokens,
            chunk_overlap_tokens=profile.chunk_overlap_tokens,
        )
        if not chunks:
            return await self._write_empty(kb_id, kb_file_id, profile.id)

        # 并发 embed
        embed_start = time.perf_counter()
        semaphore = asyncio.Semaphore(self._chunk_concurrency)

        async def embed_chunk(c: Chunk) -> tuple[Chunk, list[float]]:
            async with semaphore:
                # review chatgpt-codex P2：必须用 profile.model_id 而非全局默认 embedding 路由。
                # 否则蓝绿 / A-B 测试时 profile 切换形同虚设（向量仍用 ai_task_routing.embedding）。
                vec = await self.llm.embed(
                    c.text,
                    embedding_model_id=profile.model_id,
                    strict_embedding_model_id=True,
                )
                return c, vec

        try:
            embed_results = await asyncio.gather(*(embed_chunk(c) for c in chunks))
        except Exception as exc:
            embed_ms = (time.perf_counter() - embed_start) * 1000
            logger.warning(
                "kb_indexer.embed_failed",
                extra={"data": {
                    "kb_id": kb_id, "kb_file_id": kb_file_id,
                    "chunks": len(chunks), "embed_ms": round(embed_ms, 2),
                    "error": str(exc)[:200],
                }},
            )
            return KBIndexOutcome(kb_file_id, profile.id, 0, doc_chars, 0, "FAILED",
                                  f"embedding 失败: {type(exc).__name__}: {str(exc)[:300]}")
        embed_ms = (time.perf_counter() - embed_start) * 1000

        first_dim = len(embed_results[0][1]) if embed_results else 0
        if first_dim <= 0:
            return KBIndexOutcome(kb_file_id, profile.id, 0, doc_chars, 0, "FAILED", "embedding 返回空向量")
        for c, vec in embed_results:
            if len(vec) != first_dim:
                return KBIndexOutcome(kb_file_id, profile.id, 0, doc_chars, 0, "FAILED",
                                      f"chunk #{c.index} dim={len(vec)} 与首个 chunk dim={first_dim} 不一致")

        # 估算 token 数 = chunks 累加
        doc_tokens = sum(c.tokens for c, _ in embed_results)

        # 单事务：DELETE 旧向量 + INSERT 新 chunks。蓝绿留 Phase2，这里直接写 active。
        db_start = time.perf_counter()
        try:
            async with self.pool.acquire() as conn:
                async with conn.transaction():
                    await conn.execute(
                        "DELETE FROM kb_embeddings WHERE kb_file_id = $1 AND profile_id = $2",
                        kb_file_id, profile.id,
                    )
                    rows_to_insert = [
                        (
                            kb_file_id, kb_id, profile.id,
                            c.index, c.text, c.parent_text,
                            vec, first_dim,
                            target_status, c.tokens,
                        )
                        for c, vec in embed_results
                    ]
                    await conn.executemany(
                        """
                        INSERT INTO kb_embeddings
                            (kb_file_id, kb_id, profile_id,
                             chunk_index, chunk_text, parent_text,
                             embedding, embedding_dim,
                             status, token_count)
                        VALUES ($1,$2,$3,$4,$5,$6,$7,$8,$9,$10)
                        """,
                        rows_to_insert,
                    )
        except Exception as exc:
            db_ms = (time.perf_counter() - db_start) * 1000
            logger.warning(
                "kb_indexer.db_write_failed",
                extra={"data": {
                    "kb_id": kb_id, "kb_file_id": kb_file_id, "profile_id": profile.id,
                    "chunks": len(chunks), "db_ms": round(db_ms, 2), "error": str(exc)[:200],
                }},
            )
            return KBIndexOutcome(kb_file_id, profile.id, 0, doc_chars, doc_tokens, "FAILED",
                                  f"写入向量库失败: {type(exc).__name__}: {str(exc)[:300]}")

        db_ms = (time.perf_counter() - db_start) * 1000
        logger.info(
            "kb_indexer.ok",
            extra={"data": {
                "kb_id": kb_id, "kb_file_id": kb_file_id, "profile_id": profile.id,
                "chunks": len(chunks), "dim": first_dim,
                "embed_ms": round(embed_ms, 2), "db_ms": round(db_ms, 2),
            }},
        )
        return KBIndexOutcome(
            kb_file_id=kb_file_id,
            profile_id=profile.id,
            chunk_count=len(embed_results),
            doc_chars=doc_chars,
            doc_tokens=doc_tokens,
            status="SUCCEEDED",
        )

    async def _write_empty(self, kb_id: int, kb_file_id: int, profile_id: int) -> KBIndexOutcome:
        """空文档：仅删除目标 profile 下该 file 的旧向量，标 SUCCEEDED chunk_count=0。

        review chatgpt-codex P1 修复：之前 DELETE 全表（不限 profile_id）会把
        蓝绿场景下当前 active 的 embeddings 一并清掉，迁移若失败则文件对 active
        profile 不可搜。改为按 (kb_file_id, profile_id) 双键限定。
        """
        async with self.pool.acquire() as conn:
            await conn.execute(
                "DELETE FROM kb_embeddings WHERE kb_file_id = $1 AND profile_id = $2",
                kb_file_id, profile_id,
            )
        return KBIndexOutcome(kb_file_id, profile_id, 0, 0, 0, "SUCCEEDED")
